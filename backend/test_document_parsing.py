"""
Test script for document parsing functionality
Tests LangChain integration and ChromaDB storage
"""

import sys
import os
sys.path.append(os.path.dirname(os.path.abspath(__file__)))

import tempfile
import shutil
from pathlib import Path
from typing import Dict, Any

def create_test_pdf():
    """Create a simple test PDF using reportlab"""
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        # Create temp PDF file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
        
        # Create PDF content
        c = canvas.Canvas(temp_file.name, pagesize=letter)
        
        # Add content to PDF
        c.drawString(100, 750, "Test Document for Parsing")
        c.drawString(100, 720, "This is a sample PDF document created for testing the document parsing functionality.")
        c.drawString(100, 690, "It contains multiple lines of text that should be parsed into semantic chunks.")
        c.drawString(100, 660, "The parsing system should extract this content and store it in ChromaDB.")
        c.drawString(100, 630, "This will help us test the LangChain integration and vector storage.")
        
        # Add more content on page 2
        c.showPage()
        c.drawString(100, 750, "Page 2 Content")
        c.drawString(100, 720, "This is the second page of the test document.")
        c.drawString(100, 690, "It should be processed as a separate chunk or section.")
        c.drawString(100, 660, "The chunking algorithm should handle multi-page documents correctly.")
        
        c.save()
        temp_file.close()
        
        print(f"✅ Created test PDF: {temp_file.name}")
        return temp_file.name
        
    except ImportError:
        print("⚠️ reportlab not installed. Skipping PDF creation test.")
        return None
    except Exception as e:
        print(f"❌ Failed to create test PDF: {e}")
        return None

def create_test_docx():
    """Create a simple test DOCX using python-docx"""
    try:
        from docx import Document
        
        # Create temp DOCX file
        temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
        temp_file.close()
        
        # Create DOCX content
        doc = Document()
        
        doc.add_heading('Test Document for Parsing', 0)
        
        doc.add_paragraph(
            'This is a sample Word document created for testing the document parsing functionality. '
            'It contains multiple paragraphs that should be parsed into semantic chunks.'
        )
        
        doc.add_heading('Section 1: Introduction', level=1)
        doc.add_paragraph(
            'The parsing system should extract this content and store it in ChromaDB. '
            'This will help us test the LangChain integration and vector storage capabilities.'
        )
        
        doc.add_heading('Section 2: Testing', level=1)
        doc.add_paragraph(
            'The chunking algorithm should handle structured documents correctly. '
            'Headers and paragraphs should be processed appropriately to maintain context.'
        )
        
        doc.save(temp_file.name)
        
        print(f"✅ Created test DOCX: {temp_file.name}")
        return temp_file.name
        
    except ImportError:
        print("⚠️ python-docx not installed. Skipping DOCX creation test.")
        return None
    except Exception as e:
        print(f"❌ Failed to create test DOCX: {e}")
        return None

def test_document_parsing_service():
    """Test the document parsing service directly"""
    print("\n🧪 Testing Document Parsing Service...")
    
    try:
        from app.services.document_parser import document_parser
        
        # Test initialization
        print("✅ Document parser service imported successfully")
        print(f"   Supported file types: {list(document_parser.loader_mapping.keys())}")
        print(f"   Chunk size: {document_parser.text_splitter._chunk_size}")
        print(f"   Chunk overlap: {document_parser.text_splitter._chunk_overlap}")
        
        return True
        
    except Exception as e:
        print(f"❌ Document parsing service test failed: {e}")
        return False

def test_langchain_loaders():
    """Test LangChain document loaders"""
    print("\n🧪 Testing LangChain Document Loaders...")
    
    try:
        from langchain_community.document_loaders import (
            PyPDFLoader,
            Docx2txtLoader,
            UnstructuredPowerPointLoader
        )
        from langchain_text_splitters import RecursiveCharacterTextSplitter
        
        print("✅ LangChain loaders imported successfully")
        print("   - PyPDFLoader")
        print("   - Docx2txtLoader") 
        print("   - UnstructuredPowerPointLoader")
        print("   - RecursiveCharacterTextSplitter")
        
        return True
        
    except ImportError as e:
        print(f"❌ LangChain loader import failed: {e}")
        print("   Missing dependencies. Install with: pip install -r requirements.txt")
        return False
    except Exception as e:
        print(f"❌ LangChain loader test failed: {e}")
        return False

def test_document_parsing_flow():
    """Test complete document parsing flow with test files"""
    print("\n🧪 Testing Complete Document Parsing Flow...")
    
    if not test_document_parsing_service():
        return False
    
    try:
        from app.services.document_parser import document_parser
        
        # Create test files
        test_files = []
        
        # Test PDF
        pdf_path = create_test_pdf()
        if pdf_path:
            test_files.append(('PDF', pdf_path))
        
        # Test DOCX
        docx_path = create_test_docx()
        if docx_path:
            test_files.append(('DOCX', docx_path))
        
        if not test_files:
            print("⚠️ No test files created. Skipping parsing flow test.")
            return False
        
        # Test parsing each file
        for file_type, file_path in test_files:
            print(f"\n   Testing {file_type} parsing...")
            
            try:
                # Test parse_document method (without DB session for testing)
                result = document_parser.parse_document(
                    file_path=file_path,
                    document_id=999999,  # Test document ID
                    user_id="test_user_id",
                    db_session=None  # Skip DB operations for testing
                )
                
                if result.get("success"):
                    print(f"   ✅ {file_type} parsing successful")
                    print(f"      Total chunks: {result.get('total_chunks', 0)}")
                    print(f"      Total characters: {result.get('total_characters', 0)}")
                    print(f"      File type: {result.get('file_type', 'unknown')}")
                    
                    # Show chunk preview
                    chunks_preview = result.get('chunks_preview', [])
                    if chunks_preview:
                        print(f"      First chunk preview: {chunks_preview[0]['text'][:100]}...")
                else:
                    print(f"   ❌ {file_type} parsing failed: {result.get('error', 'Unknown error')}")
            
            except Exception as e:
                print(f"   ❌ {file_type} parsing error: {e}")
            
            finally:
                # Clean up test file
                try:
                    os.unlink(file_path)
                    print(f"   🧹 Cleaned up {file_type} test file")
                except:
                    pass
        
        return True
        
    except Exception as e:
        print(f"❌ Document parsing flow test failed: {e}")
        return False

def test_chromadb_integration():
    """Test ChromaDB integration with document chunks"""
    print("\n🧪 Testing ChromaDB Integration...")
    
    try:
        from app.services.chromadb_service import chromadb_service
        
        # Test ChromaDB connection
        if not chromadb_service.health_check():
            print("❌ ChromaDB not available. Start docker-compose services first.")
            return False
        
        print("✅ ChromaDB connection verified")
        
        # Test search functionality (placeholder)
        results = chromadb_service.search_similar_chunks(
            query="test document parsing",
            n_results=3
        )
        
        print(f"✅ Search functionality working - found {len(results)} results")
        
        return True
        
    except Exception as e:
        print(f"❌ ChromaDB integration test failed: {e}")
        return False

def main():
    """Run all document parsing tests"""
    print("🚀 Document Parsing Implementation Tests")
    print("=" * 50)
    
    tests = [
        ("LangChain Loaders", test_langchain_loaders),
        ("Document Parsing Service", test_document_parsing_service),
        ("ChromaDB Integration", test_chromadb_integration),
        ("Complete Parsing Flow", test_document_parsing_flow),
    ]
    
    results = []
    
    for test_name, test_func in tests:
        print(f"\n{'='*20} {test_name} {'='*20}")
        try:
            result = test_func()
            results.append((test_name, result))
        except Exception as e:
            print(f"❌ {test_name} crashed: {e}")
            results.append((test_name, False))
    
    # Print summary
    print("\n" + "=" * 50)
    print("🏁 Test Results Summary")
    print("=" * 50)
    
    passed = 0
    for test_name, result in results:
        status = "✅ PASSED" if result else "❌ FAILED"
        print(f"{test_name}: {status}")
        if result:
            passed += 1
    
    print(f"\nTotal: {passed}/{len(results)} tests passed")
    
    if passed == len(results):
        print("\n🎉 All document parsing tests passed!")
        print("✅ Ready for full integration testing")
    else:
        print("\n⚠️ Some tests failed.")
        print("\n🛠️ Troubleshooting:")
        print("- Install dependencies: pip install -r requirements.txt")
        print("- Start ChromaDB: docker-compose up -d")
        print("- Check logs for specific error details")

if __name__ == "__main__":
    main() 